"""Generate IronMen Halfway House program documents."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

OUTPUT_DIR = Path(__file__).parent
ORG = "IronMen Halfway House"
TAGLINE = "Forging God-fearing, responsible, and successful men"
TODAY = date.today().strftime("%B %d, %Y")


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for level in range(1, 4):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "Arial"
        heading.font.color.rgb = RGBColor(0, 0, 0)


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IRONMEN")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(166, 124, 46)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ORG)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle).italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Version 1.0 | {TODAY}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TAGLINE).italic = True

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def create_intake_criteria() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(
        doc,
        "Intake Criteria & Minimum Requirements",
        "Eligibility standards for the 90-day residential program",
    )

    add_heading(doc, "1. Purpose")
    doc.add_paragraph(
        "These criteria ensure that every man admitted to the IronMen Halfway House is ready, "
        "eligible, and likely to benefit from the 90-day residential program. Admission is selective. "
        "Meeting minimum requirements does not guarantee acceptance."
    )

    add_heading(doc, "2. Non-Negotiable Requirements")
    add_table(
        doc,
        ["#", "Requirement", "Verification"],
        [
            ["1", "Male, aged 18 years or older", "National ID / passport"],
            ["2", "Willing to live on-site for the full 90 days (12 weeks)", "Signed program agreement"],
            ["3", "Committed to a God-fearing, faith-based lifestyle during residence", "Pastoral interview"],
            ["4", "Willing to abstain from pornography, gambling, drugs, and alcohol for entire stay", "Written commitment + random checks"],
            ["5", "Medically stable — no condition requiring hospitalisation during program", "Medical clearance form"],
            ["6", "Mentally stable — not in acute psychiatric crisis", "Counsellor / clinician assessment"],
            ["7", "No active violent behaviour or threat to other residents or staff", "Background check + interview"],
            ["8", "Literate enough to follow training, journaling, and house rules (Kiswahili or English)", "Basic literacy screening"],
            ["9", "Genuinely motivated to change — not attending only due to external pressure", "Director interview"],
            ["10", "Agrees to all house rules, daily schedule, and accountability measures", "Signed handbook acknowledgement"],
        ],
    )

    add_heading(doc, "3. Preferred Qualities (Strengthen Application)")
    add_bullets(
        doc,
        [
            "Acknowledges personal struggles openly (addiction, unemployment, lack of direction)",
            "Has at least one positive reference (pastor, employer, family member, community leader)",
            "Willing to be placed in employment or apprenticeship after graduation",
            "Demonstrates respect for authority and willingness to receive correction",
            "No pending criminal matters that would require leaving the program",
            "Has completed or is willing to complete detoxification if substance history exists (case by case)",
        ],
    )

    add_heading(doc, "4. Automatic Disqualifiers")
    add_bullets(
        doc,
        [
            "Under 18 years of age",
            "Unwilling to participate in faith-based activities (devotion, prayer, Scripture study)",
            "Active, unrepentant involvement in pornography, gambling, or substance use at intake",
            "History of sexual offences or violence against women/children",
            "Severe untreated mental illness posing risk to self or others",
            "Refusal to submit to house rules, curfew, phone policy, or drug testing",
            "Known gang activity or criminal enterprise involvement",
            "Medical condition requiring specialised care unavailable at the facility",
            "Dismissive or hostile attitude toward mentorship and correction",
        ],
    )

    add_heading(doc, "5. Intake Process (5 Steps)")
    add_numbered(
        doc,
        [
            "Initial enquiry — applicant or referral partner contacts IronMen",
            "Pre-screening — age, motivation, and basic eligibility confirmed by phone or in person",
            "Application & documentation — completed intake form, ID copy, references, medical clearance",
            "Interview panel — House Director + Mentor + Pastor/Counsellor (when available)",
            "Admission decision — admit, waitlist, or decline with written reason",
        ],
    )

    add_heading(doc, "6. Required Documents at Intake")
    add_bullets(
        doc,
        [
            "Completed Resident Intake Form",
            "Copy of National ID or passport",
            "Two reference letters (one spiritual, one character/professional)",
            "Medical clearance certificate (within 30 days)",
            "Signed House Rules Handbook acknowledgement",
            "Signed Program Agreement (90-day commitment)",
            "Emergency contact details",
            "Photograph for resident file",
        ],
    )

    add_heading(doc, "7. Scoring Rubric (Interview Panel)")
    doc.add_paragraph("Each panel member scores the applicant 1–5 on the following. Minimum average score of 3.5 required for admission.")
    add_table(
        doc,
        ["Area", "1 (Poor)", "3 (Acceptable)", "5 (Excellent)"],
        [
            ["Motivation to change", "Forced attendance", "Willing but uncertain", "Deeply committed"],
            ["Faith openness", "Hostile to faith", "Respectful, open", "Actively seeking God"],
            ["Discipline potential", "Defiant", "Can follow rules with support", "Self-disciplined"],
            ["Honesty", "Deceptive answers", "Mostly honest", "Transparent about struggles"],
            ["Work readiness", "Refuses work", "Open to training", "Eager for apprenticeship"],
            ["Community fit", "Disruptive tendencies", "Neutral", "Humble, teachable spirit"],
        ],
    )

    add_heading(doc, "8. Capacity & Waiting List")
    doc.add_paragraph(
        "The house accommodates a maximum of 10 residents per cohort. When full, qualified applicants "
        "are placed on a waiting list for the next intake. A new cohort begins after the previous group graduates."
    )

    add_heading(doc, "9. Approval Authority")
    doc.add_paragraph(
        "Final admission authority rests with the House Director, in consultation with at least one mentor "
        "and one pastoral adviser. All admission decisions are documented in the resident file."
    )

    path = OUTPUT_DIR / "01_Intake_Criteria_and_Requirements.docx"
    doc.save(path)
    return path


def create_house_rules() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc, "House Rules Handbook", "Standards of living for all residents")

    add_heading(doc, "1. Vision & Foundation")
    doc.add_paragraph(
        "IronMen Halfway House exists to transform young men into God-fearing, responsible, and successful "
        "adults. Every rule serves that purpose. Rules are not punishment — they are the structure within "
        "which men are forged."
    )
    doc.add_paragraph("Foundation Scripture: Proverbs 27:17 — As iron sharpens iron, so one man sharpens another.")

    add_heading(doc, "2. Core Principles")
    add_bullets(
        doc,
        [
            "Fear of the Lord — honour God in speech, conduct, and attitude",
            "Radical honesty — no lying, hiding, or deception among brothers",
            "Mutual respect — every man is made in the image of God",
            "Personal responsibility — own your actions, chores, and growth",
            "Brotherhood — we rise together; no man is left behind",
        ],
    )

    add_heading(doc, "3. Daily Schedule (Mandatory)")
    doc.add_paragraph("All residents follow the structured daily schedule without exception unless authorised by the House Director.")
    add_table(
        doc,
        ["Time", "Activity"],
        [
            ["05:30", "Wake up, personal devotion"],
            ["06:00", "Group fitness / exercise"],
            ["07:00", "Shower, grooming, hygiene inspection"],
            ["07:45", "Breakfast and assigned chores"],
            ["08:30", "Morning training / workshop"],
            ["12:00", "Lunch and rest"],
            ["13:30", "Afternoon skills session"],
            ["16:00", "Work readiness / apprenticeship activity"],
            ["17:30", "Structured free time (reading, journaling)"],
            ["18:30", "Dinner"],
            ["19:30", "Evening devotion and accountability circle"],
            ["21:00", "Lights-out preparation"],
            ["22:00", "Lights out — silence in dormitory areas"],
        ],
    )
    doc.add_paragraph(
        "Friday Exception: On Fridays, residents join the IronMen Brotherhood Service at 19:00 (1 hour) "
        "instead of the regular 19:30 accountability circle. Lights-out schedule remains unchanged."
    )

    add_heading(doc, "4. Friday Brotherhood Service (Weekly)")
    doc.add_paragraph(
        "Every Friday evening, all IronMen brothers gather for a 1-hour church-style service — worship, "
        "teaching, testimony, and fellowship. Residents are required to attend. Alumni, mentors, and "
        "brotherhood members are strongly encouraged to attend."
    )
    add_bullets(
        doc,
        [
            "Time: Friday 19:00 – 20:00 (7:00 PM – 8:00 PM), 60 minutes",
            "Phase 1: Online via Zoom until a permanent venue is secured",
            "Phase 2: In-person at designated church, hall, or IronMen facility",
            "Dress code: Clean, modest, presentable attire — this is worship",
            "Arrive 5 minutes early (Zoom) or 10 minutes early (in-person)",
            "No phones during service except for Zoom participation when online",
            "Attendance recorded every week",
        ],
    )

    add_heading(doc, "5. Conduct & Character")
    add_numbered(
        doc,
        [
            "Speak respectfully to staff, mentors, and fellow residents at all times.",
            "No fighting, bullying, intimidation, or threats — zero tolerance.",
            "No stealing, borrowing without permission, or misuse of another man's property.",
            "Confess relapses immediately to your mentor — hidden sin grows in darkness.",
            "Complete all assigned chores and training tasks to the required standard.",
            "Attend all devotionals, accountability circles, and scheduled sessions.",
            "Maintain a teachable spirit — receive correction without argument.",
        ],
    )

    add_heading(doc, "6. Purity & Vice Policy (Zero Tolerance)")
    doc.add_paragraph("The following are strictly prohibited on premises and during the entire 90-day program:")
    add_bullets(
        doc,
        [
            "Pornography — no explicit material, suggestive content, or lustful conduct",
            "Gambling — no betting, casinos, sports gambling, or games of chance for money",
            "Drugs & alcohol — no possession, use, or intoxication; random testing may apply",
            "Inappropriate relationships — no romantic or sexual activity on premises",
            "Profanity and vulgar speech — especially in anger",
        ],
    )
    doc.add_paragraph(
        "First confirmed violation: immediate meeting with Director, written warning, increased accountability. "
        "Second violation: probation with mentor supervision. Third violation or serious breach: dismissal from program."
    )

    add_heading(doc, "7. Hygiene & Grooming Standards")
    add_bullets(
        doc,
        [
            "Shower daily; use deodorant and maintain oral hygiene",
            "Hair neat, beard trimmed or clean-shaven (resident choice, but must be tidy)",
            "Nails clean and trimmed",
            "Laundry done on assigned days — no dirty clothing in living areas",
            "Bed made every morning before breakfast",
            "Personal space kept clean during weekly inspection",
            "Dress modestly — no sagging trousers, offensive slogans, or inappropriate attire",
            "Sunday and graduation events: formal clean attire as directed",
        ],
    )

    add_heading(doc, "8. Room & Property")
    add_bullets(
        doc,
        [
            "Each resident is assigned a bed/room — no swapping without approval",
            "No food storage in rooms except sealed snacks approved by staff",
            "No candles, incense, lighters, or hazardous items",
            "House property must be returned in good condition",
            "Damage caused intentionally must be repaired or paid for by the resident",
        ],
    )

    add_heading(doc, "9. Curfew & Leave Policy")
    add_bullets(
        doc,
        [
            "All residents must be on premises by 20:00 unless on approved work/apprenticeship activity",
            "No overnight leave during the 90-day program except family emergency approved by Director",
            "Day leave requires written permission, mentor awareness, and return by stated time",
            "Unauthorised absence is grounds for immediate dismissal",
        ],
    )

    add_heading(doc, "10. Visitors & Communication")
    add_bullets(
        doc,
        [
            "Visitors allowed only on designated days (e.g. Sundays 14:00–17:00) with Director approval",
            "No visitors of the opposite sex in private areas",
            "Phones surrendered at intake or restricted per phone policy (see below)",
            "Phone use: 30 minutes daily at 17:30 unless Director grants work-related exception",
            "No social media accounts that access explicit or gambling content",
            "All calls may be made in designated area with accountability",
        ],
    )

    add_heading(doc, "11. Phone & Internet Policy")
    add_bullets(
        doc,
        [
            "Option A (recommended): Phone stored by staff; issued only during approved phone time",
            "Option B: Smartphone with accountability software installed and monitored",
            "No private internet browsing — house computer for job search and training only",
            "Violation of phone policy treated as potential purity breach",
        ],
    )

    add_heading(doc, "12. Financial Rules")
    add_bullets(
        doc,
        [
            "Residents surrender gambling/betting apps and accounts at intake",
            "Personal money held in trust if misuse history exists — released with mentor guidance",
            "Mandatory participation in budgeting and stewardship training",
            "No lending or borrowing money between residents without mentor approval",
            "Theft of any amount results in immediate dismissal and possible legal action",
        ],
    )

    add_heading(doc, "13. Health & Safety")
    add_bullets(
        doc,
        [
            "Report illness to staff immediately",
            "No self-medication without staff approval",
            "Fire drills and safety briefings are mandatory",
            "Smoking/vaping prohibited on premises",
            "Weapons of any kind strictly forbidden",
        ],
    )

    add_heading(doc, "14. Conflict Resolution")
    add_numbered(
        doc,
        [
            "Address the issue directly and respectfully with the person involved.",
            "If unresolved, bring the matter to your mentor within 24 hours.",
            "Mentor facilitates a meeting between parties.",
            "If still unresolved, House Director makes a final binding decision.",
            "No gossip, slander, or talking behind a brother's back.",
        ],
    )

    add_heading(doc, "15. Graduation Requirements")
    doc.add_paragraph("To graduate from the 90-day program, a resident must:")
    add_bullets(
        doc,
        [
            "Complete all 12 weeks of curriculum with documented attendance",
            "Maintain vice-free conduct for final 30 days of program",
            "Achieve minimum skill average of 5/10 across all 8 pillars",
            "Pass final hygiene, discipline, and work-readiness inspection",
            "Have employment or apprenticeship placement secured (or active pipeline)",
            "Receive recommendation from assigned mentor and House Director",
            "Participate in graduation ceremony and alumni brotherhood onboarding",
        ],
    )

    add_heading(doc, "16. Acknowledgement")
    doc.add_paragraph(
        "I, ________________________________ (print name), have read and understood the IronMen Halfway House "
        "Rules Handbook. I agree to abide by these rules for the full 90-day duration of my residence."
    )
    doc.add_paragraph("Signature: _________________________   Date: _________________________")
    doc.add_paragraph("Witness (Director/Mentor): _________________________   Date: _________________________")

    path = OUTPUT_DIR / "02_House_Rules_Handbook.docx"
    doc.save(path)
    return path


def create_intake_form() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc, "Resident Intake Form", "Confidential — for admission assessment only")

    add_heading(doc, "Section A: Personal Information")
    fields_a = [
        "Full Name:", "Date of Birth:", "Age:", "National ID / Passport No.:",
        "Phone Number:", "Email (if any):", "Home Address:", "County / Region:",
        "Marital Status:", "Number of Dependents:", "Languages Spoken:",
        "Highest Education Level:", "Current Employment Status:",
    ]
    for field in fields_a:
        doc.add_paragraph(f"{field} _________________________________________________")

    add_heading(doc, "Section B: Emergency Contact")
    for field in ["Name:", "Relationship:", "Phone:", "Address:"]:
        doc.add_paragraph(f"{field} _________________________________________________")

    add_heading(doc, "Section C: Program Motivation")
    doc.add_paragraph("Why do you want to join the IronMen Halfway House? (Write in your own words)")
    doc.add_paragraph("_" * 90)
    doc.add_paragraph("_" * 90)
    doc.add_paragraph("What do you hope to achieve after 90 days?")
    doc.add_paragraph("_" * 90)
    doc.add_paragraph("Who referred you to IronMen? _________________________________________________")

    add_heading(doc, "Section D: Vice & Struggle Disclosure (Honesty Required)")
    doc.add_paragraph("Check all that apply and provide brief details:")
    add_table(
        doc,
        ["Struggle", "Yes/No", "Duration", "Currently Active?", "Notes"],
        [
            ["Pornography", "", "", "", ""],
            ["Gambling / Betting", "", "", "", ""],
            ["Alcohol", "", "", "", ""],
            ["Drug use", "", "", "", ""],
            ["Anger / violence", "", "", "", ""],
            ["Unemployment", "", "", "", ""],
            ["Debt / financial mismanagement", "", "", "", ""],
            ["Other", "", "", "", ""],
        ],
    )

    add_heading(doc, "Section E: Medical & Mental Health")
    add_table(
        doc,
        ["Question", "Yes/No", "Details"],
        [
            ["Any chronic illness?", "", ""],
            ["Currently on medication?", "", ""],
            ["History of hospitalisation?", "", ""],
            ["History of self-harm?", "", ""],
            ["Receiving counselling/therapy?", "", ""],
            ["Allergies?", "", ""],
            ["Fit for physical exercise?", "", ""],
        ],
    )
    doc.add_paragraph("Medical clearance certificate attached: Yes / No")

    add_heading(doc, "Section F: Legal & Background")
    add_table(
        doc,
        ["Question", "Yes/No", "Details"],
        [
            ["Any pending court cases?", "", ""],
            ["Ever convicted of a crime?", "", ""],
            ["Ever charged with violence?", "", ""],
            ["On probation or parole?", "", ""],
        ],
    )

    add_heading(doc, "Section G: References (Minimum 2 Required)")
    doc.add_paragraph("Reference 1 — Spiritual (Pastor / Church Leader)")
    for field in ["Name:", "Role:", "Phone:", "Years known:", "Character summary:"]:
        doc.add_paragraph(f"{field} _________________________________________________")

    doc.add_paragraph("Reference 2 — Character / Professional")
    for field in ["Name:", "Relationship:", "Phone:", "Years known:", "Character summary:"]:
        doc.add_paragraph(f"{field} _________________________________________________")

    add_heading(doc, "Section H: Intake Criteria Checklist (Staff Use)")
    add_table(
        doc,
        ["Criterion", "Met?", "Notes"],
        [
            ["Male, 18+", "", ""],
            ["90-day commitment", "", ""],
            ["Faith-based participation", "", ""],
            ["Vice-free commitment", "", ""],
            ["Medical clearance", "", ""],
            ["Mental stability", "", ""],
            ["No violence risk", "", ""],
            ["Literacy adequate", "", ""],
            ["Motivated to change", "", ""],
            ["House rules accepted", "", ""],
        ],
    )

    add_heading(doc, "Section I: Interview Panel Scores")
    add_table(
        doc,
        ["Area", "Director (1-5)", "Mentor (1-5)", "Pastor (1-5)", "Average"],
        [
            ["Motivation", "", "", "", ""],
            ["Faith openness", "", "", "", ""],
            ["Discipline potential", "", "", "", ""],
            ["Honesty", "", "", "", ""],
            ["Work readiness", "", "", "", ""],
            ["Community fit", "", "", "", ""],
        ],
    )
    doc.add_paragraph("Decision: ADMIT / WAITLIST / DECLINE")
    doc.add_paragraph("Reason: _________________________________________________")

    add_heading(doc, "Section J: Agreements & Signatures")
    doc.add_paragraph("I declare that the information provided is true to the best of my knowledge.")
    doc.add_paragraph("Applicant Signature: _________________________  Date: _________________________")
    doc.add_paragraph("House Director: _________________________  Date: _________________________")
    doc.add_paragraph("Assigned Room: _______   Cohort: _________________________")

    path = OUTPUT_DIR / "03_Resident_Intake_Form.docx"
    doc.save(path)
    return path


def create_mentor_guide() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc, "Mentor Guide", "Roles, responsibilities, and standards for IronMen mentors")

    add_heading(doc, "1. The Mentor's Calling")
    doc.add_paragraph(
        "A mentor at IronMen is not merely a supervisor. You are a brother who walks alongside a younger man, "
        "sharpening him as iron sharpens iron. Your role is to model godly manhood, speak truth in love, and "
        "hold residents accountable to the standards of the house and the Word of God."
    )

    add_heading(doc, "2. Mentor Qualifications")
    add_bullets(
        doc,
        [
            "Male, aged 25 or older (preferably 30+)",
            "Active, mature Christian faith — member of a local church",
            "Demonstrated integrity in work, family, and finances",
            "Free from active addiction struggles (minimum 2 years clean if history exists)",
            "Emotionally stable and able to receive training",
            "Pass background check",
            "Committed to mentor 1–3 residents for full 90-day cohort",
            "Available for minimum 6 hours per week (structured + informal)",
        ],
    )

    add_heading(doc, "3. Core Responsibilities")
    add_table(
        doc,
        ["Responsibility", "Frequency", "Details"],
        [
            ["1-on-1 mentoring", "Weekly (45 min)", "Review struggles, goals, Scripture, accountability"],
            ["Training session delivery", "As scheduled", "Lead or co-lead assigned curriculum modules"],
            ["Skill assessment", "Every 4 weeks", "Score all 8 pillars; document in app/records"],
            ["Daily visibility", "Daily", "Observe conduct, punctuality, hygiene, attitude"],
            ["Accountability circle", "Daily (evening)", "Facilitate or participate in group honesty time"],
            ["Incident reporting", "As needed", "Report rule violations to Director within 1 hour"],
            ["Placement support", "Weeks 8–12", "Help resident prepare for job/apprenticeship"],
            ["Graduation recommendation", "Week 12", "Submit written recommendation"],
        ],
    )

    add_heading(doc, "4. Weekly Mentor Rhythm")
    add_table(
        doc,
        ["Day", "Activity", "Time"],
        [
            ["Monday", "Review weekend conduct; set weekly goals with mentees", "08:30 or 1-on-1"],
            ["Tuesday", "Co-lead skills workshop", "08:30 or 13:30"],
            ["Wednesday", "1-on-1 mentoring sessions", "Flexible"],
            ["Thursday", "Observe work-readiness activity; give feedback", "16:00"],
            ["Friday", "Submit weekly mentor report to Director", "By 17:00"],
            ["Saturday", "Community service or fitness (optional attendance)", "As scheduled"],
            ["Sunday", "Church attendance with residents (strongly encouraged)", "Morning"],
        ],
    )

    add_heading(doc, "5. What to Cover in 1-on-1 Sessions")
    add_numbered(
        doc,
        [
            "Open with prayer.",
            "Ask: How was your walk with God this week?",
            "Review vice accountability — any slips? Confession and plan.",
            "Review goals from last week — completed or not? Why?",
            "Hygiene/grooming and discipline observations.",
            "Financial stewardship check (weeks 4+).",
            "Work readiness progress (weeks 8+).",
            "Assign one Scripture memory verse and one challenge for the week.",
            "Close with prayer and encouragement.",
        ],
    )

    add_heading(doc, "6. Documentation Requirements")
    doc.add_paragraph("All mentors must document the following in the IronMen app or resident file:")
    add_bullets(
        doc,
        [
            "Every 1-on-1 session — date, topics covered, concerns, action items",
            "Every training session led — module, attendance, observations",
            "Skill assessments at weeks 4, 8, and 12",
            "Any incident or relapse — within 1 hour of knowledge",
            "Weekly mentor summary — submitted every Friday",
            "Final graduation recommendation at week 12",
        ],
    )

    add_heading(doc, "7. Red Flags — Report Immediately")
    add_bullets(
        doc,
        [
            "Suspected drug/alcohol use or gambling relapse",
            "Accessing pornography or inappropriate content",
            "Threats, fights, or bullying",
            "Self-harm statements or severe depression",
            "Unauthorised absence or curfew violation",
            "Dishonesty in accountability circle",
            "Theft or misuse of property",
            "Romantic/sexual contact on premises",
            "Resistance to authority escalating to defiance",
        ],
    )

    add_heading(doc, "8. How to Correct a Resident")
    add_numbered(
        doc,
        [
            "Correct privately first — never humiliate publicly unless safety requires it.",
            "Be specific: name the behaviour, not the person's character.",
            "Reference the house rule or Scripture standard.",
            "State the expected change clearly.",
            "Document the conversation.",
            "Follow up within 48 hours.",
            "Escalate to Director if no change after two private corrections.",
        ],
    )

    add_heading(doc, "9. Boundaries Mentors Must Keep")
    add_bullets(
        doc,
        [
            "No lending money to residents",
            "No private meetings behind closed doors without visibility",
            "No sharing personal addiction history in graphic detail",
            "No favouritism — treat all mentees fairly",
            "No romantic or inappropriate relationships with residents",
            "No confidentiality for matters involving safety or illegal activity",
            "No undermining the House Director's authority",
        ],
    )

    add_heading(doc, "10. Mentor Weekly Report Template")
    doc.add_paragraph("Mentor Name: _______________  Week: ___  Cohort: _______________")
    add_table(
        doc,
        ["Resident", "Attendance", "Vice Status", "Skill Progress", "Concerns", "Action Items"],
        [
            ["", "", "", "", "", ""],
            ["", "", "", "", "", ""],
            ["", "", "", "", "", ""],
        ],
    )
    doc.add_paragraph("Submitted to Director: Yes / No    Date: _______________")

    add_heading(doc, "11. Mentor Covenant")
    doc.add_paragraph(
        "I, ________________________________ (print name), commit to serving as a mentor at IronMen Halfway House "
        "for the duration of cohort _______________. I will walk with integrity, speak truth in love, document "
        "faithfully, and prioritise the transformation of the men entrusted to my care."
    )
    doc.add_paragraph("Signature: _________________________   Date: _________________________")

    path = OUTPUT_DIR / "04_Mentor_Guide.docx"
    doc.save(path)
    return path


def create_budget() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Monthly Budget"

    header_fill = PatternFill("solid", fgColor="D4A853")
    input_fill = PatternFill("solid", fgColor="FFF2CC")
    total_fill = PatternFill("solid", fgColor="E2EFDA")
    header_font = Font(name="Arial", bold=True, color="000000")
    bold = Font(name="Arial", bold=True)
    normal = Font(name="Arial")
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws["A1"] = "IronMen Halfway House — Monthly Operating Budget"
    ws["A1"].font = Font(name="Arial", bold=True, size=14)
    ws["A2"] = f"Capacity: 10 residents | 90-day cohorts | Currency: KES | Prepared: {TODAY}"
    ws["A2"].font = Font(name="Arial", italic=True, size=10)

    headers = ["Category", "Line Item", "Monthly (KES)", "Notes / Assumption"]
    for col, value in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=value)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    rows = [
        ("Facility", "Rent / Lease", 120000, "3-4 bedroom house or converted facility, Nairobi satellite town"),
        ("Facility", "Electricity", 15000, "10 residents + staff; solar reduces long-term"),
        ("Facility", "Water & Sewer", 5000, "Municipal supply"),
        ("Facility", "Internet & WiFi", 5000, "Training, job search, communication"),
        ("Facility", "Maintenance & Repairs", 10000, "Plumbing, paint, furniture, contingency"),
        ("Facility", "Insurance (property/liability)", 8000, "Essential for safeguarding organisation"),
        ("Food", "Groceries (10 residents)", 90000, "3 meals/day; KES 300/day per resident"),
        ("Food", "Cooking gas & supplies", 8000, "Gas, oil, spices, utensils replacement"),
        ("Staff", "House Director stipend", 50000, "May be volunteer; budget for sustainability"),
        ("Staff", "Mentors (2 x stipend)", 60000, "KES 30,000 each; may be volunteer"),
        ("Staff", "Security (night watch)", 20000, "Optional but recommended"),
        ("Staff", "Part-time cook/cleaner", 25000, "Rotating chores reduce need; backup support"),
        ("Program", "Training materials & books", 8000, "Bibles, workbooks, hygiene kits"),
        ("Program", "Fitness equipment upkeep", 5000, "Mats, weights, sports gear"),
        ("Program", "Transport (site visits, church)", 15000, "Matatu/group transport for outings"),
        ("Program", "Medical / first aid", 5000, "First aid, basic meds, health emergencies"),
        ("Program", "Drug testing kits", 3000, "Random screening"),
        ("Program", "Graduation & intake events", 5000, "Certificates, meals, ceremony costs"),
        ("Admin", "Office supplies & printing", 4000, "Forms, reports, documentation"),
        ("Admin", "Communication (airtime)", 3000, "Staff coordination"),
        ("Admin", "Legal & registration fees", 5000, "Amortised monthly; NGO/CBO compliance"),
        ("Admin", "Bank charges & misc", 2000, "Transaction fees"),
        ("Contingency", "Emergency reserve (10%)", 0, "Auto-calculated below total"),
    ]

    start_row = 5
    for i, (cat, item, amount, note) in enumerate(rows):
        r = start_row + i
        ws.cell(r, 1, cat).font = normal
        ws.cell(r, 2, item).font = normal
        c = ws.cell(r, 3, amount)
        c.font = Font(name="Arial", color="0000FF")
        c.fill = input_fill
        c.number_format = '#,##0'
        c.border = border
        ws.cell(r, 4, note).font = normal
        for col in (1, 2, 4):
            ws.cell(r, col).border = border

    subtotal_row = start_row + len(rows)
    ws.cell(subtotal_row, 2, "SUBTOTAL").font = bold
    ws.cell(subtotal_row, 3, f"=SUM(C{start_row}:C{subtotal_row - 1})").font = bold
    ws.cell(subtotal_row, 3).number_format = '#,##0'
    ws.cell(subtotal_row, 3).fill = total_fill

    contingency_row = subtotal_row + 1
    ws.cell(contingency_row, 2, "Contingency (10%)").font = bold
    ws.cell(contingency_row, 3, f"=ROUND(C{subtotal_row}*0.1,0)").font = bold
    ws.cell(contingency_row, 3).number_format = '#,##0'
    ws.cell(contingency_row, 3).fill = total_fill

    total_row = contingency_row + 1
    ws.cell(total_row, 2, "TOTAL MONTHLY OPERATING COST").font = Font(name="Arial", bold=True, size=12)
    ws.cell(total_row, 3, f"=C{subtotal_row}+C{contingency_row}").font = Font(name="Arial", bold=True, size=12)
    ws.cell(total_row, 3).number_format = '#,##0'
    ws.cell(total_row, 3).fill = PatternFill("solid", fgColor="C6EFCE")

    per_resident_row = total_row + 2
    ws.cell(per_resident_row, 2, "Cost per resident per month").font = bold
    ws.cell(per_resident_row, 3, f"=C{total_row}/10").font = bold
    ws.cell(per_resident_row, 3).number_format = '#,##0'

    per_cohort_row = per_resident_row + 1
    ws.cell(per_cohort_row, 2, "Cost per 90-day cohort (3 months)").font = bold
    ws.cell(per_cohort_row, 3, f"=C{total_row}*3").font = Font(name="Arial", bold=True, size=12)
    ws.cell(per_cohort_row, 3).number_format = '#,##0'
    ws.cell(per_cohort_row, 3).fill = PatternFill("solid", fgColor="D9E1F2")

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 34
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 48

    # Annual projection sheet
    annual = wb.create_sheet("Annual Projection")
    annual["A1"] = "Annual Projection (4 Cohorts / Year)"
    annual["A1"].font = Font(name="Arial", bold=True, size=14)
    annual["A3"] = "Item"
    annual["B3"] = "Amount (KES)"
    annual["A4"] = "Monthly operating cost"
    annual["B4"] = "='Monthly Budget'!C" + str(total_row)
    annual["A5"] = "Annual operating cost (12 months)"
    annual["B5"] = "=B4*12"
    annual["A6"] = "Cost per cohort (90 days)"
    annual["B6"] = "='Monthly Budget'!C" + str(per_cohort_row)
    annual["A7"] = "Annual cohort costs (4 cohorts)"
    annual["B7"] = "=B6*4"
    annual["A9"] = "Suggested annual fundraising target"
    annual["B9"] = "=B5*1.15"
    annual["A10"] = "Suggested cost recovery per resident (optional fee)"
    annual["B10"] = "='Monthly Budget'!C" + str(per_resident_row)

    for r in range(3, 11):
        annual.cell(r, 1).font = normal if r > 3 else bold
        annual.cell(r, 2).number_format = '#,##0'
        if r >= 5:
            annual.cell(r, 2).font = bold

    annual.column_dimensions["A"].width = 40
    annual.column_dimensions["B"].width = 20

    # Startup costs sheet
    startup = wb.create_sheet("Startup Costs")
    startup["A1"] = "One-Time Startup Costs (Estimate)"
    startup["A1"].font = Font(name="Arial", bold=True, size=14)
    startup_headers = ["Item", "Cost (KES)", "Notes"]
    for col, val in enumerate(startup_headers, 1):
        startup.cell(3, col, val).font = header_font
        startup.cell(3, col).fill = header_fill

    startup_items = [
        ("Lease deposit (2 months)", 240000, "Refundable"),
        ("Furniture (10 beds, mattresses, lockers)", 200000, "Durable, simple"),
        ("Kitchen setup (fridge, cooker, utensils)", 80000, ""),
        ("Training room setup (chairs, whiteboard, projector)", 60000, ""),
        ("Fitness equipment", 50000, "Basic weights, mats, balls"),
        ("Hygiene starter kits (10 residents)", 30000, "Towels, toiletries, grooming"),
        ("Legal registration (NGO/CBO)", 50000, "One-time"),
        ("Initial food & supplies stock", 50000, "First month buffer"),
        ("Branding & signage", 20000, ""),
        ("Technology (laptop, printer, app setup)", 80000, ""),
    ]
    for i, (item, cost, note) in enumerate(startup_items, 4):
        startup.cell(i, 1, item)
        startup.cell(i, 2, cost).number_format = '#,##0'
        startup.cell(i, 2).font = Font(name="Arial", color="0000FF")
        startup.cell(i, 3, note)

    startup_total_row = 4 + len(startup_items)
    startup.cell(startup_total_row, 1, "TOTAL STARTUP").font = bold
    startup.cell(startup_total_row, 2, f"=SUM(B4:B{startup_total_row - 1})").font = bold
    startup.cell(startup_total_row, 2).number_format = '#,##0'

    startup.column_dimensions["A"].width = 42
    startup.column_dimensions["B"].width = 16
    startup.column_dimensions["C"].width = 28

    path = OUTPUT_DIR / "05_Operating_Budget.xlsx"
    wb.save(path)
    return path


def create_friday_service_guide() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(
        doc,
        "Friday Brotherhood Service Guide",
        "Weekly worship & fellowship — Zoom first, then in-person",
    )

    add_heading(doc, "1. Purpose")
    doc.add_paragraph(
        "The Friday Brotherhood Service is the weekly spiritual gathering of IronMen — a church-style meeting "
        "for worship, Scripture teaching, testimony, and fellowship. It unites house residents, alumni, "
        "mentors, and brotherhood members in one hour of God-centred brotherhood."
    )

    add_heading(doc, "2. Schedule")
    add_table(
        doc,
        ["Detail", "Standard"],
        [
            ["Day", "Every Friday"],
            ["Time", "19:00 – 20:00 (7:00 PM – 8:00 PM)"],
            ["Duration", "60 minutes (strict)"],
            ["Attendance", "Mandatory for residents; encouraged for all brothers"],
            ["Phase 1", "Zoom (online) — until venue secured"],
            ["Phase 2", "In-person at church, hall, or IronMen facility"],
        ],
    )

    add_heading(doc, "3. Order of Service (60 Minutes)")
    add_table(
        doc,
        ["Duration", "Item", "Led By"],
        [
            ["5 min", "Opening Prayer", "Worship Leader"],
            ["15 min", "Praise & Worship", "Music Team"],
            ["5 min", "Scripture Reading", "Assigned Brother"],
            ["25 min", "Message / Teaching", "Speaker"],
            ["5 min", "Testimony or Accountability Moment", "Open"],
            ["5 min", "Closing Prayer & Announcements", "Director / Pastor"],
        ],
    )

    add_heading(doc, "4. Zoom Phase Setup")
    add_numbered(
        doc,
        [
            "Create a recurring Zoom meeting for every Friday at 19:00.",
            "Enable waiting room for security.",
            "Assign a co-host (mentor) to admit brothers.",
            "Share link via IronMen app every Monday.",
            "Test audio 15 minutes before service.",
            "Record attendance in the app at start of service.",
            "Mute all participants on entry; unmute for testimony only.",
        ],
    )

    add_heading(doc, "5. Transition to In-Person")
    doc.add_paragraph("When moving to a physical location:")
    add_bullets(
        doc,
        [
            "Secure a venue: local church, community hall, or IronMen house common area",
            "Set up chairs facing speaker; worship area to one side",
            "Assign greeters, ushers, and security",
            "Continue using the app for attendance and announcements",
            "Maintain Zoom option for brothers who cannot travel (hybrid)",
            "Announce transition 2 weeks in advance",
        ],
    )

    add_heading(doc, "6. 12-Week Teaching Rotation")
    topics = [
        ("1", "The Fear of the Lord", "Proverbs 9:10"),
        ("2", "Breaking Free from Bondage", "John 8:36"),
        ("3", "Discipline — The Mark of a Man", "1 Corinthians 9:27"),
        ("4", "Stewardship & Money", "Luke 16:10"),
        ("5", "Purity in a Corrupt Age", "1 Thessalonians 4:3-4"),
        ("6", "Brotherhood & Accountability", "Proverbs 27:17"),
        ("7", "Work as Worship", "Colossians 3:23"),
        ("8", "Forgiveness & Restoration", "1 John 1:9"),
        ("9", "Leadership at Home & Work", "1 Timothy 3:1-5"),
        ("10", "Overcoming Gambling & Greed", "Hebrews 13:5"),
        ("11", "Preparing for Graduation", "Joshua 1:9"),
        ("12", "Commissioning & New Beginnings", "Philippians 1:6"),
    ]
    add_table(doc, ["Week", "Topic", "Scripture"], topics)

    add_heading(doc, "7. Roles & Responsibilities")
    add_table(
        doc,
        ["Role", "Responsibilities"],
        [
            ["House Director", "Overall oversight, announcements, attendance, venue/Zoom management"],
            ["Worship Leader", "Song selection, opening prayer, flow of worship"],
            ["Speaker", "25-minute message aligned to weekly topic"],
            ["Mentor / Host", "Admit Zoom participants, greet in-person, support testimony time"],
            ["Tech Brother", "Zoom audio/video, lyrics display, recording if approved"],
        ],
    )

    add_heading(doc, "8. Conduct During Service")
    add_bullets(
        doc,
        [
            "Arrive on time — lateness is disrespectful to God and brothers",
            "Dress presentably — clean shirt, no sagging, no offensive graphics",
            "No eating, phone scrolling, or side conversations during service",
            "Participate in worship — stand, sing, engage",
            "Testimonies must be edifying, honest, and under 3 minutes",
            "Confidentiality: what is shared in service stays in the brotherhood",
        ],
    )

    add_heading(doc, "9. Attendance & Accountability")
    doc.add_paragraph(
        "Attendance is recorded every Friday in the IronMen app. Residents who miss two consecutive "
        "services without approved excuse meet with the House Director. Alumni and members are encouraged "
        "but not penalised — however, consistent attendance builds the brotherhood."
    )

    add_heading(doc, "10. Zoom Meeting Template")
    doc.add_paragraph("Paste your recurring meeting details below:")
    doc.add_paragraph("Zoom Link: _________________________________________________")
    doc.add_paragraph("Meeting ID: _________________________________________________")
    doc.add_paragraph("Passcode: _________________________________________________")
    doc.add_paragraph("Host: _________________________________________________")

    path = OUTPUT_DIR / "06_Friday_Brotherhood_Service_Guide.docx"
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_intake_criteria(),
        create_house_rules(),
        create_intake_form(),
        create_mentor_guide(),
        create_budget(),
        create_friday_service_guide(),
    ]
    print("Created documents:")
    for p in created:
        print(f"  - {p}")


if __name__ == "__main__":
    main()